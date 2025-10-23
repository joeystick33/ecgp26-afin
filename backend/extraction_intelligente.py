"""
Module d'extraction intelligente utilisant l'IA pour analyser des documents financiers
Capable de scanner des documents de n'importe quelle longueur (40+ pages)
"""

import io
import re
from typing import Dict, Any, Optional, List
import pandas as pd
from openai import AsyncOpenAI
import os
import json

# Import des bibliothèques PDF
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ExtracteurIntelligent:
    """
    Extracteur intelligent qui utilise l'IA pour analyser tout le document,
    peu importe sa longueur ou son format
    """
    
    def __init__(self):
        """Initialise le client OpenAI si la clé API est disponible"""
        api_key = os.getenv("OPENAI_API_KEY")
        self.client = AsyncOpenAI(api_key=api_key) if api_key else None
        self.has_openai = api_key is not None
    
    # ============================================================================
    # EXTRACTION DU TEXTE BRUT (supporte documents longs)
    # ============================================================================
    
    @staticmethod
    def extraire_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extrait TOUT le texte d'un PDF, peu importe le nombre de pages
        """
        print(f"\n=== EXTRACTION PDF: {len(file_bytes)} bytes ===")
        
        # Essayer PyMuPDF en priorité (plus rapide pour les gros documents)
        if HAS_PYMUPDF:
            try:
                return ExtracteurIntelligent._extraire_pymupdf(file_bytes)
            except Exception as e:
                print(f"PyMuPDF error: {e}")
        
        # Fallback sur pdfplumber
        if HAS_PDFPLUMBER:
            try:
                return ExtracteurIntelligent._extraire_pdfplumber(file_bytes)
            except Exception as e:
                print(f"pdfplumber error: {e}")
        
        raise Exception("Aucune bibliothèque PDF disponible")
    
    @staticmethod
    def _extraire_pymupdf(file_bytes: bytes) -> Dict[str, Any]:
        """Extraction rapide avec PyMuPDF - optimal pour documents longs"""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        texte_complet = ""
        nb_pages = len(doc)
        
        print(f"Extraction de {nb_pages} pages avec PyMuPDF...")
        
        for page_num in range(nb_pages):
            page = doc[page_num]
            texte = page.get_text()
            texte_complet += f"\n\n=== PAGE {page_num + 1} ===\n\n{texte}"
        
        doc.close()
        
        print(f"✓ Extraction terminée: {len(texte_complet)} caractères")
        
        return {
            "texte_brut": texte_complet,
            "nb_pages": nb_pages,
            "nb_caracteres": len(texte_complet),
            "methode_utilisee": "pymupdf"
        }
    
    @staticmethod
    def _extraire_pdfplumber(file_bytes: bytes) -> Dict[str, Any]:
        """Extraction avec pdfplumber"""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            texte_complet = ""
            nb_pages = len(pdf.pages)
            
            print(f"Extraction de {nb_pages} pages avec pdfplumber...")
            
            for page_num, page in enumerate(pdf.pages):
                texte = page.extract_text()
                if texte:
                    texte_complet += f"\n\n=== PAGE {page_num + 1} ===\n\n{texte}"
            
            print(f"✓ Extraction terminée: {len(texte_complet)} caractères")
            
            return {
                "texte_brut": texte_complet,
                "nb_pages": nb_pages,
                "nb_caracteres": len(texte_complet),
                "methode_utilisee": "pdfplumber"
            }
    
    @staticmethod
    def extraire_excel(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait les données depuis Excel"""
        try:
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            
            texte_complet = ""
            for nom_feuille, df in df_dict.items():
                texte_complet += f"\n\n=== FEUILLE: {nom_feuille} ===\n\n"
                texte_complet += df.to_string()
            
            return {
                "texte_brut": texte_complet,
                "nb_feuilles": len(df_dict),
                "nb_caracteres": len(texte_complet),
                "methode_utilisee": "pandas_excel"
            }
            
        except Exception as e:
            raise Exception(f"Erreur extraction Excel: {str(e)}")
    
    @staticmethod
    def extraire_csv(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait les données depuis CSV"""
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
            texte_complet = df.to_string()
            
            return {
                "texte_brut": texte_complet,
                "nb_lignes": len(df),
                "nb_caracteres": len(texte_complet),
                "methode_utilisee": "pandas_csv"
            }
            
        except Exception as e:
            raise Exception(f"Erreur extraction CSV: {str(e)}")
    
    @staticmethod
    def extraire_word(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait les données depuis Word"""
        if not HAS_DOCX:
            raise Exception("python-docx non installé")
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            texte_complet = ""
            
            for para in doc.paragraphs:
                texte_complet += para.text + "\n"
            
            return {
                "texte_brut": texte_complet,
                "nb_paragraphes": len(doc.paragraphs),
                "nb_caracteres": len(texte_complet),
                "methode_utilisee": "python_docx"
            }
            
        except Exception as e:
            raise Exception(f"Erreur extraction Word: {str(e)}")
    
    # ============================================================================
    # ANALYSE IA INTELLIGENTE (traite tout le texte)
    # ============================================================================
    
    async def analyser_avec_ia(self, texte_complet: str) -> Dict[str, Any]:
        """
        Utilise l'IA pour extraire intelligemment les données financières
        du texte complet, peu importe la longueur
        """
        if not self.has_openai:
            print("⚠️ Pas de clé OpenAI - utilisation du fallback regex")
            return self._extraction_fallback_regex(texte_complet)
        
        print(f"\n=== ANALYSE IA: {len(texte_complet)} caractères ===")
        
        # Limiter le texte si trop long (GPT a des limites de tokens)
        # On garde les 50000 premiers caractères (environ 12500 tokens)
        texte_analyse = texte_complet[:50000]
        if len(texte_complet) > 50000:
            print(f"⚠️ Texte tronqué à 50000 caractères (original: {len(texte_complet)})")
        
        prompt = f"""Tu es un expert comptable. Analyse ce document financier et extrais TOUTES les données financières que tu trouves pour l'année N (la plus récente) ET l'année N-1 (l'année précédente).

DOCUMENT À ANALYSER:
{texte_analyse}

INSTRUCTIONS IMPORTANTES:

*** PRIORITÉ ABSOLUE 1 : IDENTIFIER L'ENTREPRISE ***
- LIS LE DOCUMENT ATTENTIVEMENT pour trouver le NOM DE L'ENTREPRISE (en général en haut du document)
- IDENTIFIE LE SECTEUR D'ACTIVITÉ en analysant :
  * Le nom de l'entreprise (ex: "HÔTEL" dans le nom = secteur hotellerie)
  * L'activité décrite dans le document
  * Les postes comptables (ex: "Ventes chambres", "Restauration" = hotellerie)
  * Le code APE/NAF si mentionné
- SECTEURS POSSIBLES : "hotellerie", "restauration", "commerce", "industrie", "services", "btp", "transport", "sante", "autre"
- Si le document mentionne "hôtel", "chambres", "hébergement" → secteur = "hotellerie"
- Si le document mentionne "restaurant", "repas" → secteur = "restauration"

*** PRIORITÉ ABSOLUE 2 : EXTRAIRE LES DONNÉES FINANCIÈRES POUR N ET N-1 ***
🚨🚨🚨 CRUCIAL - LIS BIEN CETTE INSTRUCTION 🚨🚨🚨

Les documents comptables ont TOUJOURS 2 colonnes de chiffres côte à côte :
  COLONNE GAUCHE = Année N (la plus récente : 2024, 31/12/2024)
  COLONNE DROITE = Année N-1 (l'année précédente : 2023, 31/12/2023)

EXEMPLE CONCRET de ce que tu dois voir dans le document :
┌────────────────────────────────────────────────────────┐
│                              31/12/2024    31/12/2023  │
│                                  N            N-1      │
│                                                        │
│ Chiffre d'affaires         1 200 000      980 000     │
│ Achats consommés            -400 000     -350 000     │
│ Total actif                5 000 000    4 500 000     │
└────────────────────────────────────────────────────────┘

Ce qui DOIT donner dans le JSON :
{{
  "chiffre_affaires": 1200000,        ← COLONNE GAUCHE (N)
  "chiffre_affaires_n1": 980000,      ← COLONNE DROITE (N-1)
  "achats_consommes": 400000,
  "achats_consommes_n1": 350000,
  "total_actif": 5000000,
  "total_actif_n1": 4500000,
  ...
}}

⚠️ RÈGLES IMPÉRATIVES :
1. Pour CHAQUE poste comptable, tu DOIS extraire LES DEUX valeurs (N et N-1)
2. Si tu vois 2 colonnes de chiffres → EXTRAIS LES DEUX (c'est presque toujours le cas)
3. N'OUBLIE PAS les champs _n1 : ils sont AUSSI IMPORTANTS que les champs N
4. Convertis K€ en euros (×1000), M€ en euros (×1000000)
5. Si vraiment il n'y a qu'une seule colonne (très rare) : mets 0 pour tous les _n1
6. Retourne UNIQUEMENT le JSON complet

FORMAT DE RÉPONSE (JSON strict):
{{
  "nom_entreprise": <string ou null>,
  "secteur_activite": <string: "hotellerie", "restauration", "commerce", "industrie", "services", "btp", "transport", "sante", "autre">,
  "description_activite": <string: brève description de l'activité>,
  "forme_juridique": <string: SAS, SARL, SA, EURL, SCI, etc. ou null>,
  "capital_social": <nombre en euros ou null>,
  "ville": <string: ville du siège social ou null>,
  "annee_creation": <int: année de création ou null>,
  "annee_n": <int: année N>,
  "annee_n1": <int: année N-1 ou null si non disponible>,
  
  "chiffre_affaires": <nombre ou null>,
  "autres_produits_exploitation": <nombre ou null>,
  "subventions_exploitation": <nombre ou null>,
  
  "achats_marchandises": <nombre ou null>,
  "achats_matieres_premieres": <nombre ou null>,
  "achats_consommes": <nombre ou null>,
  "autres_charges_externes": <nombre ou null>,
  
  "impots_taxes": <nombre ou null>,
  "charges_personnel": <nombre ou null>,
  "remunerations_personnel": <nombre ou null>,
  "charges_sociales": <nombre ou null>,
  
  "dotations_amortissements": <nombre ou null>,
  "dotations_provisions": <nombre ou null>,
  "autres_charges_exploitation": <nombre ou null>,
  
  "produits_financiers": <nombre ou null>,
  "charges_financieres": <nombre ou null>,
  
  "produits_exceptionnels": <nombre ou null>,
  "charges_exceptionnelles": <nombre ou null>,
  
  "impot_benefices": <nombre ou null>,
  
  "total_actif": <nombre ou null>,
  "capitaux_propres": <nombre ou null>,
  
  "immobilisations_incorporelles": <nombre ou null>,
  "immobilisations_corporelles": <nombre ou null>,
  "immobilisations_financieres": <nombre ou null>,
  
  "stocks": <nombre ou null>,
  "stocks_matieres_premieres": <nombre ou null>,
  "stocks_produits_finis": <nombre ou null>,
  "stocks_marchandises": <nombre ou null>,
  
  "creances_clients": <nombre ou null>,
  "creances_autres": <nombre ou null>,
  "disponibilites": <nombre ou null>,
  
  "dettes_financieres": <nombre ou null>,
  "emprunts_long_terme": <nombre ou null>,
  "emprunts_court_terme": <nombre ou null>,
  "dettes_fournisseurs": <nombre ou null>,
  "dettes_fiscales_sociales": <nombre ou null>,
  
  "actif_circulant": <nombre ou null>,
  "passif_circulant": <nombre ou null>,
  
  "nombre_actions": <nombre ou null>,
  "cours_action": <nombre ou null>,
  
  "chiffre_affaires_n1": <nombre ou null>,
  "autres_produits_exploitation_n1": <nombre ou null>,
  "achats_marchandises_n1": <nombre ou null>,
  "achats_matieres_premieres_n1": <nombre ou null>,
  "achats_consommes_n1": <nombre ou null>,
  "autres_charges_externes_n1": <nombre ou null>,
  "impots_taxes_n1": <nombre ou null>,
  "charges_personnel_n1": <nombre ou null>,
  "remunerations_personnel_n1": <nombre ou null>,
  "charges_sociales_n1": <nombre ou null>,
  "dotations_amortissements_n1": <nombre ou null>,
  "dotations_provisions_n1": <nombre ou null>,
  "autres_charges_exploitation_n1": <nombre ou null>,
  "produits_financiers_n1": <nombre ou null>,
  "charges_financieres_n1": <nombre ou null>,
  "produits_exceptionnels_n1": <nombre ou null>,
  "charges_exceptionnelles_n1": <nombre ou null>,
  "impot_benefices_n1": <nombre ou null>,
  "total_actif_n1": <nombre ou null>,
  "capitaux_propres_n1": <nombre ou null>,
  "immobilisations_incorporelles_n1": <nombre ou null>,
  "immobilisations_corporelles_n1": <nombre ou null>,
  "immobilisations_financieres_n1": <nombre ou null>,
  "stocks_n1": <nombre ou null>,
  "stocks_matieres_premieres_n1": <nombre ou null>,
  "stocks_produits_finis_n1": <nombre ou null>,
  "stocks_marchandises_n1": <nombre ou null>,
  "creances_clients_n1": <nombre ou null>,
  "creances_autres_n1": <nombre ou null>,
  "disponibilites_n1": <nombre ou null>,
  "dettes_financieres_n1": <nombre ou null>,
  "emprunts_long_terme_n1": <nombre ou null>,
  "emprunts_court_terme_n1": <nombre ou null>,
  "dettes_fournisseurs_n1": <nombre ou null>,
  "dettes_fiscales_sociales_n1": <nombre ou null>,
  "actif_circulant_n1": <nombre ou null>,
  "passif_circulant_n1": <nombre ou null>
}}

IMPORTANT - RÈGLES D'EXTRACTION : 
- Les champs SANS suffixe (_n1) = Année N (la plus récente)
- Les champs AVEC suffixe _n1 = Année N-1 (l'année précédente)
- Si le document a 2 colonnes de chiffres : EXTRAIS LES DEUX (c'est presque toujours le cas)
- Si tu ne trouves vraiment qu'une seule année : mets 0 pour tous les champs _n1
- Retourne UNIQUEMENT le JSON, pas de texte avant ou après

EXEMPLE de réponse attendue si le document a 2 colonnes :
{{
  "chiffre_affaires": 1200000,
  "chiffre_affaires_n1": 980000,
  "total_actif": 5000000,
  "total_actif_n1": 4500000,
  ...
}}"""

        try:
            response = await self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Tu es un expert comptable qui extrait des données financières. Tu réponds UNIQUEMENT en JSON valide."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=4000  # Augmenté pour permettre l'extraction complète N + N-1
            )
            
            contenu = response.choices[0].message.content.strip()
            print(f"✓ Réponse IA reçue: {len(contenu)} caractères")
            
            # Nettoyer le JSON (enlever markdown si présent)
            contenu = contenu.replace("```json", "").replace("```", "").strip()
            
            # Parser le JSON
            donnees = json.loads(contenu)
            
            # Convertir None en 0 pour les champs numériques
            for key in donnees:
                if donnees[key] is None:
                    donnees[key] = 0
                elif isinstance(donnees[key], (int, float)):
                    donnees[key] = float(donnees[key])
            
            # Compter les champs non-nuls (seulement les nombres)
            champs_non_nuls = len([v for v in donnees.values() if isinstance(v, (int, float)) and v != 0])
            print(f"✓ Données extraites: {champs_non_nuls} champs numériques non-nuls")
            
            # Déboguer les champs entreprise
            print(f"→ Nom entreprise: {donnees.get('nom_entreprise', 'NON TROUVÉ')}")
            print(f"→ Secteur: {donnees.get('secteur_activite', 'NON TROUVÉ')}")
            print(f"→ Description: {donnees.get('description_activite', 'NON TROUVÉ')}")
            
            # Déboguer les données N vs N-1
            champs_n = len([k for k, v in donnees.items() if not k.endswith('_n1') and isinstance(v, (int, float)) and v != 0])
            champs_n1 = len([k for k, v in donnees.items() if k.endswith('_n1') and isinstance(v, (int, float)) and v != 0])
            print(f"→ Champs année N (sans _n1): {champs_n} champs non-nuls")
            print(f"→ Champs année N-1 (avec _n1): {champs_n1} champs non-nuls")
            
            # Afficher quelques exemples N-1
            if champs_n1 > 0:
                exemples_n1 = {k: v for k, v in donnees.items() if k.endswith('_n1') and isinstance(v, (int, float)) and v != 0}
                print(f"→ Exemples N-1: {list(exemples_n1.items())[:3]}")
            else:
                print("⚠️ ATTENTION: Aucune donnée N-1 extraite !")
            
            return donnees
            
        except Exception as e:
            print(f"❌ Erreur analyse IA: {e}")
            print("→ Fallback sur regex")
            return self._extraction_fallback_regex(texte_complet)
    
    # ============================================================================
    # FALLBACK REGEX (si pas d'IA disponible)
    # ============================================================================
    
    @staticmethod
    def _extraction_fallback_regex(texte: str) -> Dict[str, float]:
        """Extraction basique par regex si l'IA n'est pas disponible"""
        print("\n=== EXTRACTION FALLBACK (REGEX) ===")
        
        texte_clean = re.sub(r'[^\x00-\x7F\u00C0-\u024F\s\-\'\.,;:()]+', ' ', texte)
        texte_clean = re.sub(r'\s+', ' ', texte_clean).lower()
        
        donnees = {}
        
        # Patterns basiques pour les données principales
        patterns = {
            "chiffre_affaires": [
                r"chiffre\s+d.?affaires.*?[:\s]+([\d\s]+)",
                r"production\s+vendue.*?[:\s]+([\d\s]+)",
            ],
            "charges_personnel": [
                r"salaires?\s+et\s+traitements?.*?[:\s]+([\d\s]+)",
                r"charges?\s+personnel.*?[:\s]+([\d\s]+)",
            ],
            "total_actif": [
                r"total\s+actif.*?[:\s]+([\d\s]+)",
                r"total\s+general.*?[:\s]+([\d\s]+)",
            ],
            "capitaux_propres": [
                r"capitaux\s+propres?.*?[:\s]+([\d\s]+)",
            ],
        }
        
        for cle, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, texte_clean)
                if match:
                    try:
                        montant_str = match.group(1).replace(" ", "")
                        donnees[cle] = float(montant_str)
                        break
                    except:
                        pass
        
        return donnees
    
    # ============================================================================
    # COMPLÉTER LES DONNÉES MANQUANTES
    # ============================================================================
    
    @staticmethod
    def completer_donnees_manquantes(donnees: Dict[str, Any]) -> Dict[str, Any]:
        """
        Complète les données manquantes avec des valeurs par défaut
        CONSERVE tous les champs extraits (y compris N-1)
        """
        champs_requis_base = {
            "chiffre_affaires": 0,
            "autres_produits_exploitation": 0,
            "achats_consommes": 0,
            "autres_charges_externes": 0,
            "impots_taxes": 0,
            "charges_personnel": 0,
            "dotations_amortissements": 0,
            "autres_charges_exploitation": 0,
            "produits_financiers": 0,
            "charges_financieres": 0,
            "produits_exceptionnels": 0,
            "charges_exceptionnelles": 0,
            "impot_benefices": 0,
            "total_actif": 0,
            "capitaux_propres": 0,
            "dettes_financieres": 0,
            "stocks": 0,
            "creances_clients": 0,
            "disponibilites": 0,
            "dettes_fournisseurs": 0,
            "actif_circulant": 0,
            "passif_circulant": 0,
            "nombre_actions": None,
            "cours_action": None,
        }
        
        # Commencer avec les champs de base
        donnees_completes = champs_requis_base.copy()
        
        # IMPORTANT: Copier TOUS les champs extraits (y compris _n1, nom_entreprise, etc.)
        for key, value in donnees.items():
            if value is not None and value != "":
                try:
                    # Si c'est un nombre, le convertir en float
                    if isinstance(value, (int, float)):
                        donnees_completes[key] = float(value)
                    # Sinon garder tel quel (strings comme nom_entreprise, secteur, etc.)
                    else:
                        donnees_completes[key] = value
                except:
                    # En cas d'erreur, garder la valeur telle quelle
                    donnees_completes[key] = value
        
        return donnees_completes
