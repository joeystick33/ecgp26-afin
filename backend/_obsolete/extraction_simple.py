"""
Module d'extraction simplifié utilisant PyMuPDF uniquement
Plus léger et fiable pour l'extraction de base
"""

import io
import re
from typing import Dict, Any, Optional
import pandas as pd

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False


class ExtracteurSimple:
    """Extracteur simplifié et robuste"""
    
    @staticmethod
    def extraire_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait les données d'un PDF"""
        
        # Essayer PyMuPDF en priorité (plus léger)
        if HAS_PYMUPDF:
            try:
                return ExtracteurSimple._extraire_pymupdf(file_bytes)
            except Exception as e:
                print(f"PyMuPDF error: {e}")
        
        # Fallback sur pdfplumber
        if HAS_PDFPLUMBER:
            try:
                return ExtracteurSimple._extraire_pdfplumber(file_bytes)
            except Exception as e:
                print(f"pdfplumber error: {e}")
        
        raise Exception("Aucune bibliothèque PDF disponible")
    
    @staticmethod
    def _extraire_pymupdf(file_bytes: bytes) -> Dict[str, Any]:
        """Extraction avec PyMuPDF"""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        texte_complet = ""
        tables_detectees = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extraire le texte
            texte = page.get_text()
            texte_complet += texte + "\n"
            
            # Détecter les tableaux via extraction de texte structuré
            blocks = page.get_text("dict")["blocks"]
            table_data = ExtracteurSimple._detecter_tableaux_depuis_blocks(blocks)
            if table_data:
                tables_detectees.extend(table_data)
        
        doc.close()
        
        # Extraire les données structurées
        donnees_structurees = ExtracteurSimple._analyser_texte_et_tables(
            texte_complet, 
            tables_detectees
        )
        
        return {
            "tables": tables_detectees,
            "texte_brut": texte_complet,
            "donnees_structurees": donnees_structurees,
            "methode_utilisee": "pymupdf"
        }
    
    @staticmethod
    def _extraire_pdfplumber(file_bytes: bytes) -> Dict[str, Any]:
        """Extraction avec pdfplumber"""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            texte_complet = ""
            tables_detectees = []
            
            for page in pdf.pages:
                texte = page.extract_text()
                if texte:
                    texte_complet += texte + "\n"
                
                tables = page.extract_tables()
                if tables:
                    tables_detectees.extend(tables)
        
        donnees_structurees = ExtracteurSimple._analyser_texte_et_tables(
            texte_complet,
            tables_detectees
        )
        
        return {
            "tables": tables_detectees,
            "texte_brut": texte_complet,
            "donnees_structurees": donnees_structurees,
            "methode_utilisee": "pdfplumber"
        }
    
    @staticmethod
    def _detecter_tableaux_depuis_blocks(blocks) -> list:
        """Détecte les tableaux depuis les blocks de texte PyMuPDF"""
        # Pour l'instant, retourne une liste vide
        # L'extraction sera faite depuis le texte brut
        return []
    
    @staticmethod
    def _nettoyer_montant(texte: str) -> Optional[float]:
        """Nettoie et convertit un montant en float"""
        if not texte or not isinstance(texte, str):
            return None
        
        # Retirer tout sauf chiffres, virgule, point, tiret, espaces
        cleaned = re.sub(r'[^\d,\.\-\s]', '', texte.strip())
        
        if not cleaned or cleaned.replace(' ', '').replace('-', '') == '':
            return None
        
        # Gérer les différents formats
        # Format: 1 234 567,89 (français avec espaces)
        if ' ' in cleaned and ',' in cleaned:
            cleaned = cleaned.replace(' ', '').replace(',', '.')
        # Format: 1234,56 (français)
        elif ',' in cleaned and '.' not in cleaned:
            cleaned = cleaned.replace(',', '.')
        # Format: 1 234 567 (sans décimales)
        elif ' ' in cleaned:
            cleaned = cleaned.replace(' ', '')
        # Format: 1.234,56 (mixte européen)
        elif '.' in cleaned and ',' in cleaned:
            # Compter les séparateurs
            if cleaned.count('.') > 1 or cleaned.rfind('.') < cleaned.rfind(','):
                # Le point est le séparateur de milliers
                cleaned = cleaned.replace('.', '').replace(',', '.')
            else:
                # Format américain inversé
                cleaned = cleaned.replace(',', '')
        
        try:
            valeur = float(cleaned)
            return abs(valeur) if valeur != 0 else None
        except (ValueError, AttributeError):
            return None
    
    @staticmethod
    def _analyser_texte_et_tables(texte: str, tables: list) -> Dict[str, float]:
        """Analyse le texte et les tables pour extraire les données"""
        donnees = {}
        
        # D'abord analyser les tables si disponibles
        if tables:
            donnees_tables = ExtracteurSimple._extraire_depuis_tables(tables)
            donnees.update(donnees_tables)
        
        # Ensuite analyser le texte
        donnees_texte = ExtracteurSimple._extraire_depuis_texte(texte)
        
        # Fusionner (priorité aux tables)
        for cle, valeur in donnees_texte.items():
            if cle not in donnees and valeur:
                donnees[cle] = valeur
        
        return donnees
    
    @staticmethod
    def _extraire_depuis_tables(tables: list) -> Dict[str, float]:
        """Extrait les données depuis les tables"""
        donnees = {}
        
        # Mapping des mots-clés
        mapping = {
            "chiffre_affaires": ["chiffre", "affaires", "ca", "ventes", "produits", "revenus"],
            "achats_consommes": ["achats", "consommés", "marchandises"],
            "autres_charges_externes": ["charges externes", "services", "extérieurs"],
            "impots_taxes": ["impôts", "taxes"],
            "charges_personnel": ["personnel", "salaires", "charges sociales"],
            "dotations_amortissements": ["dotations", "amortissements"],
            "charges_financieres": ["charges financières", "intérêts"],
            "produits_financiers": ["produits financiers"],
            "resultat_exploitation": ["résultat", "exploitation"],
            "resultat_net": ["résultat net", "bénéfice"],
            "total_actif": ["total actif", "actif total"],
            "capitaux_propres": ["capitaux propres", "fonds propres"],
            "dettes_financieres": ["dettes financières", "emprunts"],
            "stocks": ["stocks"],
            "creances_clients": ["créances", "clients"],
            "disponibilites": ["disponibilités", "trésorerie"],
            "dettes_fournisseurs": ["fournisseurs", "dettes fournisseurs"]
        }
        
        for table in tables:
            if not table or len(table) < 2:
                continue
            
            for row in table:
                if not row or len(row) < 2:
                    continue
                
                libelle_cell = str(row[0]).lower() if row[0] else ""
                
                if len(libelle_cell) < 3:
                    continue
                
                # Chercher correspondance
                for cle, keywords in mapping.items():
                    if any(kw in libelle_cell for kw in keywords):
                        # Chercher le montant (dernière colonne avec valeur)
                        for cell in reversed(row[1:]):
                            if cell:
                                montant = ExtracteurSimple._nettoyer_montant(str(cell))
                                if montant and montant > 0:
                                    donnees[cle] = montant
                                    break
                        break
        
        return donnees
    
    @staticmethod
    def _extraire_depuis_texte(texte: str) -> Dict[str, float]:
        """Extrait les données depuis le texte brut avec patterns améliorés"""
        import re
        donnees = {}
        
        # Étape 1: Nettoyer le texte des caractères parasites
        # Supprimer caractères non-latins (coréen, chinois, symboles)
        texte_clean = re.sub(r'[^\x00-\x7F\u00C0-\u024F\s\-\'\.,;:()]+', ' ', texte)
        # Normaliser espaces multiples
        texte_clean = re.sub(r'\s+', ' ', texte_clean)
        texte_clean = texte_clean.lower()
        
        print(f"\n=== TEXTE NETTOYÉ (premiers 500 caractères) ===")
        print(texte_clean[:500])
        
        # Patterns améliorés - plus flexibles et robustes
        patterns = {
            "chiffre_affaires": [
                r"chiffre\s+d.?affaires.*?net.*?[:\s]+([0-9][\s0-9]{5,})",
                r"production\s+vendue.*?\(.*?services?.*?\).*?([0-9][\s0-9]{5,})",
                r"production\s+vendue.*?[:\s]+([0-9][\s0-9]{5,})",
            ],
            "achats_consommes": [
                r"achats?\s+de\s+matieres?\s+premieres?.*?[:\s]+([0-9][\s0-9]{3,})",
                r"achats?\s+de\s+marchandises.*?[:\s]+([0-9][\s0-9]{3,})",
            ],
            "autres_charges_externes": [
                r"autres?\s+achats?\s+et\s+charges?\s+externes?.*?[:)\s]+([0-9][\s0-9]{4,})",
            ],
            "impots_taxes": [
                r"impots?,?\s+taxes?\s+et\s+versements?\s+assimiles?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "charges_personnel": [
                r"salaires?\s+et\s+traitements?.*?[:)\s]+([0-9][\s0-9]{4,})",
            ],
            "charges_sociales": [
                r"charges?\s+sociales?.*?[:)\s]+([0-9][\s0-9]{4,})",
            ],
            "dotations_amortissements": [
                r"sur\s+immobilisations?.*?dotations?\s+aux\s+amortissements?.*?[:)\s]+([0-9][\s0-9]{4,})",
                r"dotations?\s+aux\s+amortissements?.*?[:)\s]+([0-9][\s0-9]{4,})",
            ],
            "charges_financieres": [
                r"total\s+des\s+charges?\s+financieres?.*?[:)\s]+([0-9][\s0-9]{3,})",
                r"interets?\s+et\s+charges?\s+assimilees?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "produits_financiers": [
                r"total\s+des\s+produits?\s+financiers?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "resultat_exploitation": [
                r"resultat\s+d.?exploitation.*?[:)\s]+([\-]?[0-9][\s0-9]{4,})",
                r"[0-9]\s*-\s*resultat\s+d.?exploitation.*?[:)\s]+([\-]?[0-9][\s0-9]{4,})",
            ],
            "resultat_net": [
                r"resultat\s+de\s+l.?exercice.*?[:)\s]+([\-]?[0-9][\s0-9]{4,})",
                r"benefice\s+ou\s+perte.*?[:)\s]+([\-]?[0-9][\s0-9]{4,})",
                r"[0-9]\s*-\s*benefice\s+ou\s+perte.*?[:)\s]+([\-]?[0-9][\s0-9]{4,})",
            ],
            "total_actif": [
                r"total\s+general.*?[:)\s]+([0-9][\s0-9]{6,})",
                r"total\s+\(?i\s+a\s+vi\)?.*?[:)\s]+([0-9][\s0-9]{6,})",
            ],
            "capitaux_propres": [
                r"total\s*i\s*[^a-z0-9]{1,50}([0-9][\s0-9]{6,})",
                r"capitaux\s+propres?.*?[:)\s]+([0-9][\s0-9]{6,})",
            ],
            "dettes_financieres": [
                r"emprunts?\s+et\s+dettes?\s+aupres\s+des\s+et.*?credit.*?[:)\s]+([0-9][\s0-9]{4,})",
                r"emprunts?\s+et\s+dettes?\s+financieres?.*?[:)\s]+([0-9][\s0-9]{4,})",
            ],
            "stocks": [
                r"stocks?\s+et\s+en[-\s]?cours.*?[:)\s]+([0-9][\s0-9]{3,})",
                r"total\s*iii.*?stocks.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "creances_clients": [
                r"creances?\s+clients?\s+et\s+comptes?\s+rattaches?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "disponibilites": [
                r"disponibilites?.*?[:)\s]+([0-9][\s0-9]{3,})",
                r"banques?\s+caisses?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
            "dettes_fournisseurs": [
                r"fournisseurs?\s+et\s+comptes?\s+rattaches?.*?[:)\s]+([0-9][\s0-9]{3,})",
            ],
        }
        
        # Extraction avec gestion des cas spéciaux
        for cle, pattern_list in patterns.items():
            for pattern in pattern_list:
                matches = list(re.finditer(pattern, texte_clean, re.IGNORECASE))
                if matches:
                    # Prendre le premier match (généralement le bon)
                    montant_str = matches[0].group(1)
                    montant = ExtracteurSimple._nettoyer_montant(montant_str)
                    if montant and montant > 0:
                        donnees[cle] = montant
                        print(f"✅ {cle}: {montant} (trouvé avec pattern: {pattern[:50]}...)")
                        break
        
        # Cas spécial: additionner salaires + charges sociales pour charges_personnel
        if "charges_personnel" in donnees and "charges_sociales" in donnees:
            donnees["charges_personnel"] = donnees["charges_personnel"] + donnees["charges_sociales"]
            del donnees["charges_sociales"]
        
        print(f"\n=== DONNÉES EXTRAITES: {len(donnees)} champs ===")
        
        return donnees
    
    # Variable globale pour le reader EasyOCR (inité une seule fois)
    _easyocr_reader = None
    
    @staticmethod
    def _get_easyocr_reader():
        """Initialise et retourne le reader EasyOCR (singleton)"""
        if ExtracteurSimple._easyocr_reader is None:
            try:
                import easyocr
                print("  🔧 Initialisation EasyOCR (première fois, peut prendre 10-20s)...")
                ExtracteurSimple._easyocr_reader = easyocr.Reader(['fr', 'en'], gpu=False, verbose=False)
                print("  ✅ EasyOCR prêt!")
            except Exception as e:
                print(f"  ⚠️ Erreur init EasyOCR: {str(e)}")
                return None
        return ExtracteurSimple._easyocr_reader
    
    @staticmethod
    def _extraire_images_word_ocr(doc) -> str:
        """Extrait le texte des images du document Word via EasyOCR"""
        texte_images = []
        
        try:
            from PIL import Image
            import numpy as np
            
            reader = ExtracteurSimple._get_easyocr_reader()
            if reader is None:
                print("  💡 CONSEIL: Utilisez le COPIER-COLLER pour un résultat optimal")
                return ""
            
            # Parcourir toutes les relations du document pour trouver les images
            nb_images = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Récupérer l'image
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        
                        # Convertir en image PIL puis numpy
                        image = Image.open(io.BytesIO(image_bytes))
                        image_np = np.array(image)
                        
                        # OCR avec EasyOCR (français + anglais)
                        nb_images += 1
                        print(f"  🔍 OCR image {nb_images}...")
                        results = reader.readtext(image_np, detail=0)  # detail=0 pour avoir juste le texte
                        
                        if results:
                            texte_ocr = '\n'.join(results)
                            texte_images.append(texte_ocr)
                            print(f"  ✅ Image {nb_images}: {len(texte_ocr)} caractères extraits")
                        else:
                            print(f"  ⚠️ Image {nb_images}: aucun texte détecté")
                    except Exception as e:
                        print(f"  ⚠️ Erreur OCR image {nb_images}: {str(e)}")
                        continue
            
            if nb_images == 0:
                print("  🔍 Aucune image trouvée dans le document")
            
            return '\n\n'.join(texte_images)
            
        except ImportError as e:
            print(f"  ⚠️ Bibliothèque manquante: {str(e)}")
            print("  💡 CONSEIL: Utilisez le COPIER-COLLER pour un résultat optimal")
            return ""
        except Exception as e:
            print(f"  ⚠️ Erreur extraction images: {str(e)}")
            return ""
    
    @staticmethod
    def extraire_word(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait depuis Word (.doc/.docx)"""
        try:
            from docx import Document
            
            # Lire le document Word
            doc = Document(io.BytesIO(file_bytes))
            
            # Extraire tout le texte
            texte_complet = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    texte_complet.append(paragraph.text)
            
            # Extraire les tableaux
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            texte = '\n'.join(texte_complet)
            
            print(f"\nWord: {len(texte)} caractères de texte, {len(tables)} tableaux")
            
            # Tenter OCR sur les images
            print("Recherche d'images dans le document...")
            texte_images = ExtracteurSimple._extraire_images_word_ocr(doc)
            
            # Combiner texte normal + texte des images
            if texte_images:
                texte_complet_avec_images = texte + "\n" + texte_images
                print(f"✅ Total avec OCR: {len(texte_complet_avec_images)} caractères")
            else:
                texte_complet_avec_images = texte
                if len(texte) < 500:
                    print("\n⚠️ PEU DE TEXTE EXTRAIT!")
                    print("💡 CONSEIL: Ouvrez le Word, copiez tout le contenu visible")
                    print("   et utilisez l'option COPIER-COLLER (plus fiable)")
            
            # Analyser le texte et les tables
            donnees_structurees = ExtracteurSimple._analyser_texte_et_tables(texte_complet_avec_images, tables)
            
            return {
                'texte_brut': texte_complet_avec_images,
                'tables': tables,
                'donnees_structurees': donnees_structurees,
                'methode_utilisee': 'python-docx + OCR' if texte_images else 'python-docx',
                'ocr_utilise': bool(texte_images)
            }
        except Exception as e:
            print(f"Erreur extraction Word: {str(e)}")
            return {
                'texte_brut': '',
                'tables': [],
                'donnees_structurees': {},
                'methode_utilisee': 'erreur',
                'erreur': str(e)
            }
    
    @staticmethod
    def extraire_excel(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait depuis Excel"""
        df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        
        tables = []
        for df in df_dict.values():
            if not df.empty:
                tables.append(df.values.tolist())
        
        return {
            "feuilles": {nom: df.to_dict('records') for nom, df in df_dict.items()},
            "donnees_structurees": ExtracteurSimple._extraire_depuis_tables(tables)
        }
    
    @staticmethod
    def extraire_csv(file_bytes: bytes) -> Dict[str, Any]:
        """Extrait depuis CSV"""
        df = pd.read_csv(io.BytesIO(file_bytes))
        
        return {
            "donnees_brutes": df.to_dict('records'),
            "donnees_structurees": ExtracteurSimple._extraire_depuis_tables([df.values.tolist()])
        }
    
    @staticmethod
    def completer_donnees_manquantes(donnees: Dict[str, float]) -> Dict[str, float]:
        """Complète avec des valeurs par défaut"""
        complet = {
            "chiffre_affaires": 0,
            "autres_produits_exploitation": 0,
            "achats_consommes": 0,
            "autres_charges_externes": 0,
            "impots_taxes": 0,
            "charges_personnel": 0,
            "dotations_amortissements": 0,
            "autres_charges_exploitation": 0,
            "produits_financiers": 0,
            "charges_financieres": 0,
            "produits_exceptionnels": 0,
            "charges_exceptionnelles": 0,
            "impot_benefices": 0,
            "total_actif": 0,
            "capitaux_propres": 0,
            "dettes_financieres": 0,
            "stocks": 0,
            "creances_clients": 0,
            "disponibilites": 0,
            "dettes_fournisseurs": 0,
            "actif_circulant": 0,
            "passif_circulant": 0
        }
        
        complet.update(donnees)
        return complet
